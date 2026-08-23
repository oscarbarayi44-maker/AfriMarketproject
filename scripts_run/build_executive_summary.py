"""Génère docs/Resume_Executif_AfriMarket.docx (5 pages max)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIGS = PROJECT_ROOT / "figures"
OUT = PROJECT_ROOT / "docs" / "Resume_Executif_AfriMarket.docx"

NAVY = RGBColor(0x1F, 0x2D, 0x50)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def body(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def add_figure(name, caption, width=6.2):
    path = FIGS / name
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GREY


# ------------------------------------------------------------------ PAGE 1 ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Résumé exécutif — Analyse stratégique AfriMarket")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Analyse de 6 mois d'activité commerciale — Data Analyst, Direction Générale")
run.font.size = Pt(11)
run.font.color.rgb = GREY

doc.add_paragraph()
h1("Contexte et mission")
body(
    "AfriMarket est une entreprise e-commerce panafricaine opérant dans 4 catégories "
    "(Électronique, Mode, Beauté, Maison) sur 8 grandes villes d'Afrique francophone. La "
    "direction constatait des variations importantes de chiffre d'affaires, un taux de retour "
    "préoccupant, des dépenses marketing élevées et des écarts de performance selon les villes. "
    "Cette analyse répond à quatre questions stratégiques : quelle catégorie prioriser, où "
    "investir géographiquement, quel canal marketing renforcer ou réduire, et comment améliorer "
    "la rétention client."
)

h1("Méthodologie")
body(
    "Le dataset brut (10 100 commandes) contenait des anomalies significatives : 100 doublons, "
    "614 remises négatives, 632 prix négatifs, 608 quantités nulles, 799 valeurs de prix extrêmes, "
    "ainsi que des incohérences de saisie (ville « Kinshassa » au lieu de « Kinshasa », catégorie "
    "« electronique » sans accent/majuscule, statut « retournée » en minuscule). Après nettoyage "
    "et normalisation, le dataset exploitable (df_clean) contient 9 400 commandes valides. "
    "Hypothèse retenue en l'absence de coût d'achat produit dans les données : une marge brute "
    "forfaitaire de 35% du chiffre d'affaires, à ajuster si un coût matière réel est communiqué."
)

h1("Performance globale (6 mois)")
kpi_table = doc.add_table(rows=2, cols=5)
kpi_table.style = "Light Grid Accent 1"
headers = ["CA total", "Profit net estimé", "Panier moyen", "Taux d'annulation", "Taux de retour"]
values = ["2,51 M$", "754 k$", "271,9 $", "1,9 %", "8,3 %"]
for i, (hcell, vcell) in enumerate(zip(headers, values)):
    kpi_table.cell(0, i).text = hcell
    kpi_table.cell(1, i).text = vcell
    for row in (0, 1):
        for p in kpi_table.cell(row, i).paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                if row == 0:
                    r.bold = True

doc.add_paragraph()
body(
    "Sur 6 mois, AfriMarket génère 2,51 M$ de chiffre d'affaires pour un profit net estimé de "
    "754 k$ (≈30% de marge nette sur le CA non annulé). Le taux d'annulation reste faible (1,9%), "
    "mais le taux de retour (8,3%) constitue un point de vigilance direct sur la rentabilité."
)

doc.add_page_break()

# ------------------------------------------------------------------ PAGE 2 ---
h1("Analyse par catégorie — Quelle catégorie prioriser ?")
add_figure("ca_par_categorie.png", "CA par catégorie", width=3.0)
body(
    "L'Électronique porte 74,6% du CA total (1,87 M$) et 82% du profit net (615 k$) : c'est le "
    "moteur de croissance de l'entreprise. Mode et Beauté affichent un profit quasi nul rapporté à "
    "leur marge brute théorique : les coûts de livraison et marketing absorbent respectivement 57% "
    "et 88% de leur marge (contre 6% pour l'Électronique), leur panier moyen étant trop faible pour "
    "couvrir le coût fixe par commande."
)
bullet("Prioriser l'investissement sur l'Électronique (référencement, stock, publicité).")
bullet("Optimiser Mode et Beauté par le panier moyen (seuils de livraison gratuite, bundles) plutôt que par le volume.")
bullet("Surveiller le taux de retour Électronique (14,1%, ~2x la moyenne) qui érode sa rentabilité.")

h1("Analyse géographique — Où investir davantage ?")
add_figure("taux_annulation_ville.png", "Taux d'annulation par ville", width=3.0)
body(
    "Kinshasa est le marché le plus performant (752 k$ de CA, 227 k$ de profit) avec un taux "
    "d'annulation quasi nul (0,26%) : marché mature à renforcer en priorité. Douala présente une "
    "anomalie opérationnelle majeure — taux d'annulation de 12,9% contre moins de 1% partout "
    "ailleurs — qui doit être auditée (paiement local, rupture de stock, fiabilité du livreur) "
    "avant tout investissement marketing supplémentaire dans cette ville."
)

doc.add_page_break()

# ------------------------------------------------------------------ PAGE 3 ---
h1("Analyse marketing — Quel canal renforcer ou réduire ?")
add_figure("roi_par_canal.png", "ROI par canal marketing", width=3.0)
body(
    "Formule : ROI = (Revenus − Coût marketing) / Coût marketing. Le canal Email affiche un ROI "
    "de 230 pour un budget marginal (2 297 $, le plus faible des 4 canaux) : il est nettement "
    "sous-financé au regard de son efficacité. Google Ads (ROI 50) reste solide et scalable. "
    "Instagram Ads génère le plus de CA (950 k$) et la meilleure rétention (54%) mais avec un ROI "
    "plus modéré (25) du fait de son budget élevé (37 k$). Influenceur cumule le pire ROI (21,5) "
    "et la pire rétention (42%) pour un budget de 16 k$."
)
bullet("Augmenter le budget Email (ROI le plus élevé, sous-financé).")
bullet("Maintenir/renforcer modérément Google Ads.")
bullet("Conserver Instagram Ads pour le volume et la fidélisation, en optimisant le ciblage.")
bullet("Réduire ou réattribuer le budget Influenceur (pire ROI et pire rétention) vers Email/Google Ads.")

h1("Analyse clients — Comment améliorer la rétention ?")
add_figure("pareto_clients.png", "Courbe de Pareto — CA cumulé par client", width=3.0)
body(
    "La base compte 1 734 clients, dont 73,2% sont récurrents (≥2 commandes) : un socle de "
    "fidélité globalement sain. L'analyse Pareto montre que 31,9% des clients génèrent 80% du CA "
    "— une concentration modérée, plutôt rassurante. La segmentation par quartiles révèle une "
    "forte polarisation : le segment Platine (25% des clients) génère 72% du CA total, contre 1,4% "
    "pour le segment Bronze."
)
bullet("Programme VIP pour le segment Platine, afin de sécuriser le cœur du chiffre d'affaires.")
bullet("Offres d'activation ciblées pour Bronze/Argent (867 clients, 7,8% du CA) pour les faire monter en gamme.")

doc.add_page_break()

# ------------------------------------------------------------------ PAGE 4 ---
h1("5 recommandations stratégiques")

numbered(
    "Prioriser l'Électronique tout en réduisant son taux de retour (74,6% du CA et 82% du profit ; "
    "fiches produits plus précises, contrôle qualité fournisseurs, politique de retour resserrée)."
)
numbered(
    "Revoir l'économie unitaire de Mode et Beauté (coûts logistiques/marketing consommant 57% et "
    "88% de leur marge) via un seuil de livraison gratuite, des bundles produits, ou une "
    "renégociation des frais de livraison."
)
numbered(
    "Auditer en urgence l'anomalie opérationnelle de Douala (12,9% d'annulation vs <1% ailleurs) "
    "avant d'y engager davantage de budget marketing."
)
numbered(
    "Réallouer le budget marketing vers l'efficacité : renforcer Email (ROI 230) et Google Ads "
    "(ROI 50), réduire Influenceur (ROI 21,5, rétention la plus faible)."
)
numbered(
    "Structurer un programme de fidélisation différencié par segment : traitement VIP pour "
    "Platine (72% du CA), offres d'activation pour Bronze/Argent (7,8% du CA) afin d'accélérer leur "
    "montée en gamme."
)

h1("Conclusion business orientée action")
body(
    "AfriMarket dispose d'un socle commercial sain : 2,51 M$ de CA sur 6 mois, 754 k$ de profit net "
    "estimé, une base client majoritairement fidèle (73,2% de récurrence) et un taux d'annulation "
    "globalement faible (1,9%). Trois leviers d'action immédiats se dégagent : consolider "
    "l'Électronique en maîtrisant ses retours, corriger la rentabilité structurelle de Mode et "
    "Beauté par le panier moyen plutôt que par le volume, et traiter sans délai l'anomalie "
    "opérationnelle de Douala qui grève silencieusement la performance géographique. Sur le plan "
    "marketing, un simple réarbitrage de budget — moins d'Influenceur, plus d'Email et de Google "
    "Ads — peut améliorer le ROI global sans augmenter la dépense totale. Enfin, la fidélisation "
    "doit devenir différenciée par segment plutôt qu'uniforme, pour sécuriser le cœur de revenu "
    "(segment Platine) tout en développant la base moins engagée. Ces indicateurs sont pilotables "
    "au mois par mois dans le dashboard Streamlit livré avec cette analyse."
)

OUT.parent.mkdir(exist_ok=True)
doc.save(OUT)
print(f"Résumé exécutif écrit : {OUT}")
